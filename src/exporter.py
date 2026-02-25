"""
설교 자동화 시스템 - Word(.docx) 출력 모듈

파이프라인 결과를 한글 Word 문서로 변환합니다.
"""

import re
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from rich.console import Console

console = Console()


class SermonExporter:
    """설교 결과물을 Word 문서로 내보내는 클래스.

    왜 Word로 출력하는가?
    - 목사님이 강단에서 직접 출력하여 읽을 수 있어야 합니다.
    - 서식(폰트 크기, 줄 간격)이 직관적으로 보여야 합니다.
    - 편집 및 추가 수정이 용이해야 합니다.
    """

    def __init__(self, output_dir: Path) -> None:
        self.output_dir = output_dir
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self) -> None:
        """문서 기본 스타일(폰트, 마진 등)을 설정합니다."""
        # 기본 폰트를 맑은 고딕으로 설정
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "맑은 고딕"
        font.size = Pt(12)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # 줄 간격 1.5
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.5
        paragraph_format.space_after = Pt(6)

        # 페이지 마진
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # 제목 스타일
        for level in range(1, 4):
            heading_style = self.doc.styles[f"Heading {level}"]
            heading_font = heading_style.font
            heading_font.name = "맑은 고딕"
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

            if level == 1:
                heading_font.size = Pt(22)
            elif level == 2:
                heading_font.size = Pt(16)
            else:
                heading_font.size = Pt(13)

    def _add_title_page(self, title: str, bible_ref: str, date_str: str) -> None:
        """표지 페이지를 추가합니다."""
        # 빈 줄 추가 (상단 여백)
        for _ in range(6):
            self.doc.add_paragraph()

        # 제목
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # 성경 구절
        ref_para = self.doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ref_para.add_run(bible_ref)
        run.font.size = Pt(16)
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # 날짜
        self.doc.add_paragraph()
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(date_str)
        run.font.size = Pt(12)
        run.font.name = "맑은 고딕"
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # 페이지 나눔
        self.doc.add_page_break()

    def _parse_and_add_content(self, content: str) -> None:
        """마크다운 형태의 텍스트를 Word 서식으로 변환하여 추가합니다.

        정교한 마크다운 파서는 아니지만, 설교 원고의 주요 구조 요소를 처리합니다:
        - ═══ 구분선 → 무시 (표지에서 처리)
        - ─── 구분선 → 구분선 또는 Heading 2
        - ## 제목 → Heading 2
        - ### 제목 → Heading 3
        - • 또는 - 리스트 → 불릿 리스트
        - [성경 구절 "인용"] → 볼드+이탤릭
        - 일반 텍스트 → Normal 스타일
        """
        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # 빈 줄
            if not line:
                i += 1
                continue

            # ═══ 큰 구분선 (무시)
            if line.startswith("═══"):
                i += 1
                continue

            # ─── 구분선 다음 줄이 제목인 경우
            if line.startswith("───"):
                # 다음 줄을 제목으로 처리
                if i + 1 < len(lines) and lines[i + 1].strip() and not lines[i + 1].strip().startswith("───"):
                    next_line = lines[i + 1].strip()
                    # 그 다음 줄이 또 구분선이면 섹션 제목
                    if i + 2 < len(lines) and lines[i + 2].strip().startswith("───"):
                        self.doc.add_heading(next_line, level=2)
                        i += 3
                        continue
                i += 1
                continue

            # ```코드 블록``` 건너뛰기
            if line.startswith("```"):
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    # 코드 블록 내용은 일반 텍스트로 추가
                    para = self.doc.add_paragraph(lines[i])
                    para.style.font.name = "Consolas"
                    i += 1
                i += 1  # 닫는 ``` 건너뛰기
                continue

            # 마크다운 헤딩
            if line.startswith("### "):
                self.doc.add_heading(line[4:], level=3)
                i += 1
                continue
            if line.startswith("## "):
                self.doc.add_heading(line[3:], level=2)
                i += 1
                continue
            if line.startswith("# "):
                self.doc.add_heading(line[2:], level=1)
                i += 1
                continue

            # 특수 기호로 시작하는 메타 정보 (📖, 📌, ⏱️ 등)
            if line and line[0] in "📖📌⏱️📅📋🔑✅🌟🔧💡⭐🎯":
                para = self.doc.add_paragraph()
                run = para.add_run(line)
                run.font.bold = True
                run.font.size = Pt(11)
                i += 1
                continue

            # 불릿 리스트 (•, -, *)
            if line.startswith(("• ", "- ", "* ", "□ ")):
                bullet_text = line.lstrip("•-* □").strip()
                para = self.doc.add_paragraph(bullet_text, style="List Bullet")
                i += 1
                continue

            # 숫자 리스트
            if re.match(r"^\d+\.\s", line):
                list_text = re.sub(r"^\d+\.\s", "", line)
                para = self.doc.add_paragraph(list_text, style="List Number")
                i += 1
                continue

            # 일반 텍스트
            para = self.doc.add_paragraph()
            # 볼드 텍스트 처리 (**text**)
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)

            i += 1

    def export(self, results: dict[str, str], bible_range: str, sermon_date: str = "") -> Path:
        """파이프라인 결과를 Word 문서로 내보냅니다.

        Args:
            results:     파이프라인 결과 딕셔너리 (phase1~phase5)
            bible_range: 사용자가 입력한 성경 범위
            sermon_date: 설교 예정일. 예: "2026년 02월 23일"

        Returns:
            생성된 Word 파일의 경로
        """
        # 표지 날짜: 설교 예정일이 있으면 사용, 없으면 오늘 날짜
        date_str = sermon_date if sermon_date else datetime.now().strftime("%Y년 %m월 %d일")

        # 파일명 타임스탬프: 설교 예정일 기준 (예: 20260302)
        if sermon_date:
            try:
                parsed = datetime.strptime(sermon_date, "%Y년 %m월 %d일")
                timestamp = parsed.strftime("%Y %m%d")
            except ValueError:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 표지
        self._add_title_page(
            title="설교 원고",
            bible_ref=f"📖 {bible_range}",
            date_str=date_str,
        )

        # Phase 5 최종 패키지 (메인 콘텐츠)
        if "phase5" in results:
            self.doc.add_heading("최종 설교 패키지", level=1)
            self._parse_and_add_content(results["phase5"])
            self.doc.add_page_break()

        # 부록: Phase 1 (본문 선정)
        if "phase1" in results:
            self.doc.add_heading("부록 A: 본문 선정 분석", level=1)
            self._parse_and_add_content(results["phase1"])
            self.doc.add_page_break()

        # 부록: Phase 3 (피드백 보고서)
        if "phase3" in results:
            self.doc.add_heading("부록 B: 통합 피드백 보고서", level=1)
            self._parse_and_add_content(results["phase3"])

        # 저장
        filename = f"{timestamp}_설교_{bible_range.replace(' ', '_')}.docx"
        filepath = self.output_dir / filename
        self.doc.save(str(filepath))

        console.print(f"  📄 Word 파일 저장: [bold green]{filepath}[/bold green]")
        return filepath
